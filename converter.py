import os
import re
import uuid
import shutil
from datetime import datetime
from typing import Tuple, Optional
import aiofiles
from docx import Document
from docx.shared import RGBColor
import asyncio

os.makedirs("uploads", exist_ok=True)
os.makedirs("static/ads", exist_ok=True)


class UzbekConverter:

    LATIN_TO_CYRILLIC = [
        ("sh", "ш"),
        ("ch", "ч"),
        ("ng", "нг"),
        ("yo", "ё"),
        ("ya", "я"),
        ("yu", "ю"),
        ("ye", "е"),
        ("g‘", "ғ"),
        ("g'", "ғ"),
        ("o‘", "ў"),
        ("o'", "ў"),
        ("a", "а"),
        ("b", "б"),
        ("d", "д"),
        ("e", "е"),
        ("f", "ф"),
        ("g", "г"),
        ("h", "ҳ"),
        ("i", "и"),
        ("j", "ж"),
        ("k", "к"),
        ("l", "л"),
        ("m", "м"),
        ("n", "н"),
        ("o", "о"),
        ("p", "п"),
        ("q", "қ"),
        ("r", "р"),
        ("s", "с"),
        ("t", "т"),
        ("u", "у"),
        ("v", "в"),
        ("x", "х"),
        ("y", "й"),
        ("z", "з"),
        ("Sh", "Ш"),
        ("Ch", "Ч"),
        ("Ng", "Нг"),
        ("Yo", "Ё"),
        ("Ya", "Я"),
        ("Yu", "Ю"),
        ("Ye", "Е"),
        ("G‘", "Ғ"),
        ("G'", "Ғ"),
        ("O‘", "Ў"),
        ("O'", "Ў"),
        ("A", "А"),
        ("B", "Б"),
        ("D", "Д"),
        ("E", "Е"),
        ("F", "Ф"),
        ("G", "Г"),
        ("H", "Ҳ"),
        ("I", "И"),
        ("J", "Ж"),
        ("K", "К"),
        ("L", "Л"),
        ("M", "М"),
        ("N", "Н"),
        ("O", "О"),
        ("P", "П"),
        ("Q", "Қ"),
        ("R", "Р"),
        ("S", "С"),
        ("T", "Т"),
        ("U", "У"),
        ("V", "В"),
        ("X", "Х"),
        ("Y", "Й"),
        ("Z", "З"),
    ]
    
    CYRILLIC_TO_LATIN = [
        ("нг", "ng"),
        ("ў", "o'"),
        ("ғ", "g'"),
        ("ё", "yo"),
        ("я", "ya"),
        ("ю", "yu"),
        ("е", "ye"),
        ("ш", "sh"),
        ("ч", "ch"),
        ("а", "a"),
        ("б", "b"),
        ("д", "d"),
        ("е", "e"),
        ("ф", "f"),
        ("г", "g"),
        ("ҳ", "h"),
        ("и", "i"),
        ("ж", "j"),
        ("к", "k"),
        ("л", "l"),
        ("м", "m"),
        ("н", "n"),
        ("о", "o"),
        ("п", "p"),
        ("қ", "q"),
        ("р", "r"),
        ("с", "s"),
        ("т", "t"),
        ("у", "u"),
        ("в", "v"),
        ("х", "x"),
        ("й", "y"),
        ("з", "z"),
        ("Нг", "Ng"),
        ("Ў", "O'"),
        ("Ғ", "G'"),
        ("Ё", "Yo"),
        ("Я", "Ya"),
        ("Ю", "Yu"),
        ("Е", "Ye"),
        ("Ш", "Sh"),
        ("Ч", "Ch"),
        ("А", "A"),
        ("Б", "B"),
        ("Д", "D"),
        ("Е", "E"),
        ("Ф", "F"),
        ("Г", "G"),
        ("Ҳ", "H"),
        ("И", "I"),
        ("Ж", "J"),
        ("К", "K"),
        ("Л", "L"),
        ("М", "M"),
        ("Н", "N"),
        ("О", "O"),
        ("П", "P"),
        ("Қ", "Q"),
        ("Р", "R"),
        ("С", "S"),
        ("Т", "T"),
        ("У", "U"),
        ("В", "V"),
        ("Х", "X"),
        ("Й", "Y"),
        ("З", "Z"),
    ]
    
    @staticmethod
    def detect_alphabet(text: str) -> str:

        cyrillic_chars = re.findall(r'[а-яёўғҳқА-ЯЁЎҒҲҚ]', text)
        cyrillic_count = len(cyrillic_chars)

        latin_chars = re.findall(r'[a-zA-Z]', text)
        latin_count = len(latin_chars)
        
        if cyrillic_count > latin_count:
            return 'cyrillic'
        else:
            return 'latin'
    
    @staticmethod
    def latin_to_cyrillic(text: str) -> str:

        text = text.replace("'", "").replace("`", "").replace("‘", "").replace("’", "")
        
        for latin, cyrillic in UzbekConverter.LATIN_TO_CYRILLIC:
            text = text.replace(latin, cyrillic)
        
        return text
    
    @staticmethod
    def cyrillic_to_latin(text: str) -> str:

        for cyrillic, latin in UzbekConverter.CYRILLIC_TO_LATIN:
            text = text.replace(cyrillic, latin)
        
        return text
    
    @staticmethod
    def convert_text(text: str) -> Tuple[str, str]:

        if not text.strip():
            return text, "none"
        
        alphabet = UzbekConverter.detect_alphabet(text)
        
        if alphabet == 'latin':
            converted = UzbekConverter.latin_to_cyrillic(text)
            direction = "latin_to_cyrillic"
        else:
            converted = UzbekConverter.cyrillic_to_latin(text)
            direction = "cyrillic_to_latin"
        
        return converted, direction


class DocxConverter:

    @staticmethod
    async def save_uploaded_file(file_content: bytes, original_filename: str) -> str:

        file_id = str(uuid.uuid4())
        extension = os.path.splitext(original_filename)[1] or ".docx"
        filename = f"{file_id}{extension}"
        filepath = os.path.join("uploads", filename)
        
        async with aiofiles.open(filepath, "wb") as f:
            await f.write(file_content)
        
        return filepath
    
    @staticmethod
    def convert_docx_file(input_path: str, output_path: str, direction: str = "auto"):

        doc = Document(input_path)
        
        # Process all paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Convert paragraph text
                if direction == "latin_to_cyrillic":
                    converted_text = UzbekConverter.latin_to_cyrillic(paragraph.text)
                elif direction == "cyrillic_to_latin":
                    converted_text = UzbekConverter.cyrillic_to_latin(paragraph.text)
                else:

                    alphabet = UzbekConverter.detect_alphabet(paragraph.text)
                    if alphabet == 'latin':
                        converted_text = UzbekConverter.latin_to_cyrillic(paragraph.text)
                    else:
                        converted_text = UzbekConverter.cyrillic_to_latin(paragraph.text)
                
                paragraph.clear()
                run = paragraph.add_run(converted_text)
                
                if paragraph.runs and len(paragraph.runs) > 0:
                    original_run = paragraph.runs[0]
                    run.bold = original_run.bold
                    run.italic = original_run.italic
                    run.underline = original_run.underline
                    run.font.size = original_run.font.size
                    run.font.name = original_run.font.name
        
        doc.save(output_path)
    
    @staticmethod
    async def convert_docx(
        file_content: bytes, 
        original_filename: str,
        direction: str = "auto"
    ) -> Tuple[Optional[str], Optional[str], str]:

        try:

            if not original_filename.lower().endswith('.docx'):
                return None, None, "Faqat .docx fayllarni yuklash mumkin"

            input_path = await DocxConverter.save_uploaded_file(file_content, original_filename)
            
            file_id = str(uuid.uuid4())
            output_filename = f"converted_{file_id}.docx"
            output_path = os.path.join("uploads", output_filename)
            
            DocxConverter.convert_docx_file(input_path, output_path, direction)
            
            return input_path, output_path, file_id
            
        except Exception as e:
            return None, None, f"DOCX konvertatsiyada xatolik: {str(e)}"
    
    @staticmethod
    def cleanup_file(filepath: str):

        if os.path.exists(filepath):
            try:
                os.remove(filepath)
            except:
                pass


async def cleanup_old_files():

    while True:
        try:
            now = datetime.now()
            for filename in os.listdir("uploads"):
                filepath = os.path.join("uploads", filename)
                if os.path.isfile(filepath):
                    mtime = datetime.fromtimestamp(os.path.getmtime(filepath))
                    if (now - mtime).total_seconds() > 15:
                        try:
                            os.remove(filepath)
                        except:
                            pass
        except:
            pass
        
        await asyncio.sleep(10)


async def save_ad_image(file_content: bytes, filename: str) -> str:

    file_id = str(uuid.uuid4())
    extension = os.path.splitext(filename)[1] or ".png"
    new_filename = f"ad_{file_id}{extension}"
    filepath = os.path.join("static", "ads", new_filename)
    
    async with aiofiles.open(filepath, "wb") as f:
        await f.write(file_content)
    
    return f"/static/ads/{new_filename}"


def get_file_size(filepath: str) -> int:

    try:
        return os.path.getsize(filepath)
    except:
        return 0
